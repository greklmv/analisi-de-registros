from docx import Document  # type: ignore
from docx.shared import Inches, Pt  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
import matplotlib.pyplot as plt
import io
import os
import pandas as pd
import logging

def create_sparkline(data, color='#006666'):
    """Genera una imagen de mini-gráfico (sparkline) para insertar en el Word."""
    plt.figure(figsize=(1.5, 0.4))
    plt.plot(data, color=color, linewidth=1.5)
    plt.fill_between(range(len(data)), data, color=color, alpha=0.1)
    plt.axis('off')
    buf = io.BytesIO()
    plt.savefig(buf, format='png', transparent=True, bbox_inches='tight', pad_inches=0)
    plt.close()
    buf.seek(0)
    return buf

def safe_add_heading(doc, text, level):
    """Añade un encabezado de forma segura."""
    try:
        doc.add_heading(text, level=level)
    except Exception as e:
        logging.warning(f"Error añadiendo encabezado '{text}': {e}")
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)

def safe_add_paragraph(doc, text="", style=None):
    """Añade un párrafo de forma segura."""
    try:
        return doc.add_paragraph(text, style=style)
    except Exception as e:
        logging.warning(f"Error añadiendo párrafo: {e}")
        p = doc.add_paragraph()
        if style and ('List Bullet' in style or 'Bullet' in style):
            p.add_run("• ")
        if text:
            p.add_run(text)
        return p

def generate_word_report(df, kpis, project_info, chart_img=None, notes=None, op_events=None, template_path="plantilla informe registros.docx", ai_conclusions=None):
    """
    Genera el informe Word utilizando la plantilla oficial de FGC.
    Restaurado a la versión Pro v5.0 con soporte para eventos operativos i IA.
    """
    
    # Búsqueda de la plantilla
    possible_paths = [
        template_path,
        os.path.join(os.getcwd(), template_path),
        "/Users/grek/IA/analisi-de-registros/plantilla informe registros.docx"
    ]
    
    doc = None
    for p in possible_paths:
        if os.path.exists(p):
            doc = Document(p)
            break
            
    if doc is None:
        doc = Document()
        doc.add_heading("INFORME D'ANÀLISI TELEMÈTRICA", 0)

    # --- 1. ACTUALIZAR METADATOS EN TABLAS ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.upper()
                if "UT:" in text: cell.text = f"UT: {project_info.get('u', project_info.get('ut', ''))}"
                elif "ANÀLISI DE REGISTRE:" in text: cell.text = f"ANÀLISI DE REGISTRE: {notes if notes else ''}"
                elif "DATA:" in text: cell.text = f"DATA: {pd.Timestamp.now().strftime('%d/%m/%Y')}"

    # --- 2. LIMPIEZA Y LLENADO DE ANÁLISIS DEL CUERPO ---
    target_found = False
    defaults = [
        "Es realitza tota la circulació en mode ATP",
        "El maquinista no actua sobre el bolet",
        "En cap moment entra l'anti bloqueig",
        "El maquinista ultrapassa l'estació"
    ]

    for p in doc.paragraphs:
        if "Anàlisi dels registres" in p.text:
            target_found = True
            # Limpiamos el texto por defecto si existe en los párrafos siguientes
            continue
        
        if target_found and any(d.lower() in p.text.lower() for d in defaults):
            p.text = "" # Vaciamos párrafos por defecto

    if target_found:
        for p in doc.paragraphs:
            if "Anàlisi dels registres" in p.text:
                p.add_run(f"\n\n{notes if notes else '(Sense observacions addicionals)'}")
                break

    # --- 3. ANÀLISI IA (Opcional) ---
    if ai_conclusions:
        safe_add_heading(doc, "Anàlisi Intel·ligència Artificial", level=1)
        safe_add_paragraph(doc, ai_conclusions)

    # --- 4. LÍNEA DE TIEMPO OPERATIVA (Eventos) ---
    if op_events:
        safe_add_heading(doc, "Línia de Temps Operativa", level=1)
        for ev in op_events:
            p = safe_add_paragraph(doc, style='List Bullet')
            t_str = ev.get('time', '--:--')
            e_txt = ev.get('event', 'Event')
            run = p.add_run(f"[{t_str}] {e_txt}")
            run.bold = True
            p.add_run(f": {ev.get('details', '')}")
            
            # Zoom Plot para anomalías
            if ev.get('is_anomaly') or "Canvi" in e_txt:
                try:
                    time_col = df.columns[0]
                    # Encontrar índice aproximado para el gráfico de zoom
                    idx_match = (df[time_col].astype(str).str.contains(t_str)).idxmax()
                    z_df = df.iloc[max(0, idx_match-20):min(len(df), idx_match+20)]
                    plt.figure(figsize=(5, 1.8))
                    plt.plot(range(len(z_df)), z_df[df.columns[1]], color='#006666')
                    plt.axvline(x=len(z_df)//2, color='red', linestyle='--')
                    plt.axis('off')
                    z_buf = io.BytesIO()
                    plt.savefig(z_buf, format='png', bbox_inches='tight')
                    plt.close()
                    doc.add_picture(z_buf, width=Inches(3.5))
                except Exception as e:
                    logging.warning(f"Error renderizando zoom para {t_str}: {e}")

    # --- 5. GRÁFICO GENERAL ---
    if chart_img:
        safe_add_heading(doc, "Visualització Telemetria General", level=1)
        doc.add_picture(io.BytesIO(chart_img), width=Inches(5.8))

    # --- 6. TABLA DE BLOQUES / KPIS ---
    if kpis and len(doc.tables) > 1:
        tbl = doc.tables[1]
        # Limpiamos filas viejas si la tabla ya tiene contenido de ejemplo
        while len(tbl.rows) > 1:
            row_to_remove = tbl.rows[-1]
            tbl._tbl.remove(row_to_remove._tr)

        for row_data in kpis[:20]: # Límite para evitar documentos infinitos
            r = tbl.add_row().cells
            if len(r) >= 5:
                r[0].text = str(row_data.get('start_time','--'))
                r[1].text = str(row_data.get('ut_indicator','---'))
                r[2].text = str(row_data.get('distance','0'))
                r[3].text = str(row_data.get('max_speed','0'))
                r[4].text = str(row_data.get('avg_speed','0'))
                
                # Sparkline
                history = row_data.get('speed_history', [])
                if len(history) > 2 and len(r) >= 6:
                    s_buf = create_sparkline(history)
                    r[5].paragraphs[0].add_run().add_picture(s_buf, width=Inches(1.0))
                
                # Anomalías
                if len(r) >= 7:
                    r[6].text = row_data.get('anomalies','')

    doc.add_paragraph(f"\nGenerat automàticament OTMR PRO v5.0 i IA - {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
